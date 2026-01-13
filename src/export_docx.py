from __future__ import annotations
from io import BytesIO
from typing import Dict
import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from .generator import Exam

def _add_heading(doc: Document, text: str, size: int = 14, bold: bool = True, center: bool = True):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if center:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def export_exam_docx(
    exam: Exam,
    spec_df: pd.DataFrame,
    mapping_report: Dict,
    validation_report: Dict,
    school_name: str = "",
    exam_title: str = "ĐỀ KIỂM TRA ĐỊNH KỲ",
    time_limit: str = "Thời gian: 40 phút",
) -> bytes:
    doc = Document()

    if school_name:
        _add_heading(doc, school_name, size=12, bold=True, center=True)
    _add_heading(doc, exam_title, size=14, bold=True, center=True)
    _add_heading(doc, f"{exam.subject} – Lớp {exam.grade} – {exam.term_or_exam}", size=12, bold=False, center=True)
    if time_limit:
        p = doc.add_paragraph(time_limit)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("I. PHẦN CÂU HỎI")
    for q in exam.questions:
        p = doc.add_paragraph()
        p.add_run(f"Câu {q.q_no} ({q.points} điểm): ").bold = True
        p.add_run(q.stem)
        if q.options:
            for opt in q.options:
                doc.add_paragraph(opt)

    doc.add_page_break()
    doc.add_paragraph("II. ĐÁP ÁN / HƯỚNG DẪN CHẤM")
    for q in exam.questions:
        p = doc.add_paragraph()
        p.add_run(f"Câu {q.q_no}: ").bold = True
        p.add_run(f"Đáp án: {q.answer}.")
        if q.explanation:
            doc.add_paragraph(q.explanation)
        if q.rubric:
            doc.add_paragraph(q.rubric)

    doc.add_page_break()
    doc.add_paragraph("III. BẢNG ĐẶC TẢ CHUẨN HÓA (Spec)")
    max_rows = min(len(spec_df), 200)
    max_cols = min(spec_df.shape[1], 11)
    table = doc.add_table(rows=max_rows + 1, cols=max_cols)
    table.style = "Table Grid"
    cols = list(spec_df.columns)[:max_cols]
    for j, c in enumerate(cols):
        table.cell(0, j).text = c
    for i in range(max_rows):
        row = spec_df.iloc[i]
        for j, c in enumerate(cols):
            table.cell(i + 1, j).text = str(row[c])

    doc.add_page_break()
    doc.add_paragraph("IV. BÁO CÁO ĐỐI SOÁT & TIN CẬY")
    doc.add_paragraph(f"Confidence: {validation_report.get('confidence', 0.0)}")
    doc.add_paragraph(f"Tổng số câu: {validation_report.get('final_total_questions', '')}")
    doc.add_paragraph(f"Tổng điểm: {validation_report.get('final_total_points', '')} / {validation_report.get('total_points_target', '')}")

    doc.add_paragraph("Tóm tắt suy diễn (global):")
    doc.add_paragraph(str(mapping_report.get("global_inferred", {})))

    doc.add_paragraph("Cảnh báo/khuyến nghị:")
    for w in validation_report.get("warnings", []):
        doc.add_paragraph(f"- {w}")

    bio = BytesIO()
    doc.save(bio)
    return bio.getvalue()
