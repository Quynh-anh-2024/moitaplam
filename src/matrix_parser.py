from __future__ import annotations
from io import BytesIO
from typing import List, Tuple
import pandas as pd

def _clean_df(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy().fillna("")
    df = df.applymap(lambda x: str(x).strip())
    df = df.loc[~(df.eq("").all(axis=1))]
    df = df.loc[:, ~(df.eq("").all(axis=0))]
    df.reset_index(drop=True, inplace=True)
    return df

def parse_matrix_file(filename: str, file_bytes: bytes) -> Tuple[List[pd.DataFrame], str]:
    name = filename.lower().strip()
    if name.endswith((".xlsx", ".xls")):
        return _parse_excel(file_bytes)
    if name.endswith(".docx"):
        return _parse_docx(file_bytes)
    if name.endswith(".pdf"):
        return _parse_pdf(file_bytes)
    raise ValueError("Unsupported file type")

def _parse_excel(file_bytes: bytes) -> Tuple[List[pd.DataFrame], str]:
    bio = BytesIO(file_bytes)
    xls = pd.ExcelFile(bio)
    tables, meta = [], []
    for sheet in xls.sheet_names:
        bio.seek(0)
        df = pd.read_excel(bio, sheet_name=sheet, header=None, engine="openpyxl")
        df = _clean_df(df)
        if len(df) == 0:
            continue
        tables.append(df)
        meta.append(f"[Excel sheet] {sheet}: rows={df.shape[0]} cols={df.shape[1]}")
    return tables, "\n".join(meta)

def _parse_docx(file_bytes: bytes) -> Tuple[List[pd.DataFrame], str]:
    from docx import Document
    bio = BytesIO(file_bytes)
    doc = Document(bio)
    paras = [p.text.strip() for p in doc.paragraphs if p.text and p.text.strip()]
    meta = "\n".join(paras[:200])

    tables = []
    for t in doc.tables:
        rows = [[c.text.strip() for c in r.cells] for r in t.rows]
        if not rows:
            continue
        df = _clean_df(pd.DataFrame(rows))
        if len(df) == 0:
            continue
        tables.append(df)
    return tables, meta

def _parse_pdf(file_bytes: bytes) -> Tuple[List[pd.DataFrame], str]:
    import pdfplumber
    bio = BytesIO(file_bytes)
    tables, meta = [], []
    with pdfplumber.open(bio) as pdf:
        for pi, page in enumerate(pdf.pages, start=1):
            text = (page.extract_text() or "").strip()
            if text:
                meta.append(f"[PDF page {pi}]\n{text}")
            try:
                extracted = page.extract_tables() or []
            except Exception:
                extracted = []
            for tbl in extracted:
                if not tbl:
                    continue
                df = _clean_df(pd.DataFrame(tbl))
                if len(df) == 0:
                    continue
                tables.append(df)
    return tables, "\n\n".join(meta)
